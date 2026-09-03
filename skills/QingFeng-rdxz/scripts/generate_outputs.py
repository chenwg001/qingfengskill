#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_outputs.py — 教育热点文章生成器 产物脚本

把 article.md（正文内用独占一行的 [[ILLU1]] [[ILLU2]] [[ILLU3]] 标记插图位置）
转为三份产物：
  - article.docx         纯文本版（无图，标记转为「（配图*）」）
  - article-images.docx  图片版（封面置顶 + 插图插在标记处）
  - article.html         HTML版（封面在正文前 + 3 张插图在对应内容后）

用法：
  python generate_outputs.py --md 篇目/article.md \
      --cover cover_toutiao.png \
      --illu illu1.png illu2.png illu3.png \
      --out 篇目/

依赖：pip install python-docx
"""
import sys
import io
import os
import re
import argparse
from pathlib import Path

# Windows GBK 兼容：统一 UTF-8 输出
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")


# ------------------------- Markdown 解析 -------------------------
def parse_blocks(md_text):
    """返回 block 列表：('h1'|'h2'|'h3'|'p'|'illu'|'hr', content)"""
    blocks = []
    lines = md_text.split("\n")
    buf = []

    def flush():
        if buf:
            blocks.append(("p", "\n".join(buf).strip()))
        buf.clear()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        m = re.match(r"^#\s+(.*)$", line)
        if m:
            flush(); blocks.append(("h1", m.group(1).strip())); i += 1; continue
        m = re.match(r"^##\s+(.*)$", line)
        if m:
            flush(); blocks.append(("h2", m.group(1).strip())); i += 1; continue
        m = re.match(r"^###\s+(.*)$", line)
        if m:
            flush(); blocks.append(("h3", m.group(1).strip())); i += 1; continue
        m = re.match(r"^\[\[ILLU(\d+)\]\]$", line.strip())
        if m:
            flush(); blocks.append(("illu", int(m.group(1)))); i += 1; continue
        if re.match(r"^---+\s*$", line) or re.match(r"^\*\*\*+\s*$", line):
            flush(); blocks.append(("hr", "")); i += 1; continue
        if line.strip() == "":
            flush(); i += 1; continue
        buf.append(line)
        i += 1
    flush()
    return blocks


# ------------------------- 文本写入工具 -------------------------
def add_runs(paragraph, text):
    """把含 **bold** 的文本写入 paragraph，支持加粗。"""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def md_inline_to_html(text):
    text = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", text)
    return text


# ------------------------- docx 生成 -------------------------
def build_docx_text(blocks, out_path):
    from docx import Document
    doc = Document()
    for kind, content in blocks:
        if kind == "illu":
            p = doc.add_paragraph()
            p.add_run(f"（配图{content}）").italic = True
        elif kind == "hr":
            doc.add_paragraph("—" * 12)
        elif kind == "h1":
            doc.add_heading(content, level=0)
        elif kind == "h2":
            doc.add_heading(content, level=1)
        elif kind == "h3":
            doc.add_heading(content, level=2)
        else:
            p = doc.add_paragraph()
            add_runs(p, content)
    doc.save(str(out_path))


def build_docx_images(blocks, out_path, cover_path, illu_paths):
    from docx import Document
    from docx.shared import Inches, Pt
    doc = Document()
    if cover_path and Path(cover_path).exists():
        doc.add_picture(str(cover_path), width=Inches(6.0))
        doc.add_paragraph()
    for kind, content in blocks:
        if kind == "illu":
            idx = content
            ip = illu_paths[idx - 1] if 1 <= idx <= len(illu_paths) else None
            if ip and Path(ip).exists():
                doc.add_picture(str(ip), width=Inches(5.5))
                cap = doc.add_paragraph()
                cap.alignment = 1
                r = cap.add_run(f"图{idx}")
                r.italic = True
                r.font.size = Pt(9)
            else:
                p = doc.add_paragraph()
                p.add_run(f"（配图{idx}）").italic = True
        elif kind == "hr":
            doc.add_paragraph("—" * 12)
        elif kind == "h1":
            doc.add_heading(content, level=0)
        elif kind == "h2":
            doc.add_heading(content, level=1)
        elif kind == "h3":
            doc.add_heading(content, level=2)
        else:
            p = doc.add_paragraph()
            add_runs(p, content)
    doc.save(str(out_path))


# ------------------------- HTML 生成 -------------------------
def build_html(blocks, out_path, cover_path, illu_paths, title, out_dir):
    out_dir = Path(out_dir).resolve()

    def rel(img):
        if img and Path(img).exists():
            try:
                rp = os.path.relpath(str(Path(img).resolve()), str(out_dir))
                return rp.replace("\\", "/")
            except Exception:
                return str(img)
        return ""

    parts = [
        "<!DOCTYPE html>",
        '<html lang="zh-CN"><head><meta charset="utf-8">',
        f"<title>{title}</title>",
        '<meta name="viewport" content="width=device-width, initial-scale=1.0">',
        """<style>
      body{font-family:-apple-system,"PingFang SC","Microsoft YaHei",sans-serif;max-width:760px;margin:0 auto;padding:24px;line-height:1.9;color:#222;}
      h1{font-size:26px;border-bottom:3px solid #c0392b;padding-bottom:8px;}
      h2{font-size:21px;margin-top:32px;color:#1a1a1a;}
      h3{font-size:18px;color:#333;}
      img.cover{width:100%;border-radius:8px;margin:8px 0 24px;}
      img.illu{width:100%;border-radius:8px;margin:20px 0 6px;}
      .cap{color:#888;font-size:13px;text-align:center;margin:0 0 18px;}
      hr{border:none;border-top:1px solid #eee;margin:28px 0;}
      p{margin:14px 0;}
    </style>""",
        "</head><body>",
    ]
    cover_rel = rel(cover_path)
    if cover_rel:
        parts.append(f'<img class="cover" src="{cover_rel}" alt="封面">')
    for kind, content in blocks:
        if kind == "illu":
            idx = content
            ip = illu_paths[idx - 1] if 1 <= idx <= len(illu_paths) else None
            r = rel(ip)
            if r:
                parts.append(f'<img class="illu" src="{r}" alt="图{idx}">')
                parts.append(f'<p class="cap">图{idx}</p>')
            else:
                parts.append(f'<p class="cap">（配图{idx}）</p>')
        elif kind == "hr":
            parts.append("<hr>")
        elif kind == "h1":
            parts.append(f"<h1>{md_inline_to_html(content)}</h1>")
        elif kind == "h2":
            parts.append(f"<h2>{md_inline_to_html(content)}</h2>")
        elif kind == "h3":
            parts.append(f"<h3>{md_inline_to_html(content)}</h3>")
        else:
            parts.append(f"<p>{md_inline_to_html(content)}</p>")
    parts.append("</body></html>")
    Path(out_path).write_text("\n".join(parts), encoding="utf-8")


# ------------------------- 主流程 -------------------------
def main():
    ap = argparse.ArgumentParser(description="教育热点文章产物生成（docx+html）")
    ap.add_argument("--md", required=True, help="article.md 路径")
    ap.add_argument("--cover", help="封面图路径（可省略则不插图）")
    ap.add_argument("--illu", nargs="*", default=[], help="插图路径列表 illu1 illu2 illu3")
    ap.add_argument("--out", required=True, help="产物输出目录（平台文件夹）")
    args = ap.parse_args()

    md_path = Path(args.md)
    if not md_path.exists():
        print(f"错误：找不到 {md_path}")
        sys.exit(1)

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    md_text = md_path.read_text(encoding="utf-8")
    blocks = parse_blocks(md_text)

    title = ""
    for kind, content in blocks:
        if kind == "h1":
            title = content
            break
    if not title:
        title = md_path.stem

    # docx 纯文本版
    docx_text = out_dir / "article.docx"
    build_docx_text(blocks, docx_text)

    # docx 图片版
    docx_img = out_dir / "article-images.docx"
    build_docx_images(blocks, docx_img, args.cover, args.illu)

    # HTML 版
    html_path = out_dir / "article.html"
    build_html(blocks, html_path, args.cover, args.illu, title, out_dir)

    print(f"已完成：{out_dir}")
    print(f"  - {docx_text.name}（纯文本版）")
    print(f"  - {docx_img.name}（图片版）")
    print(f"  - {html_path.name}（HTML版，标题：{title}）")


if __name__ == "__main__":
    main()
